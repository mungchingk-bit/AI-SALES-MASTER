"""分享工具模块 — 支持复制文本、导出文件、生成图片、生成链接。"""
import io
import os
import re
import tempfile
import threading
import time
import uuid
from http.server import HTTPServer, BaseHTTPRequestHandler

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont


def format_report_markdown(report, chat_history=None) -> str:
    """将汇报 + 话术讨论格式化为 Markdown 文本。"""
    lines = [f"# {report.source_title}\n"]
    lines.append(f"**销售**：{report.sales_name}  |  **日期**：{report.created_at[:10]}\n")
    lines.append(f"## 整体评价\n{report.summary}\n")

    if report.highlights:
        lines.append("## 做得好的地方")
        for i, h in enumerate(report.highlights, 1):
            lines.append(f"{i}. {h}")
        lines.append("")

    if report.improvements:
        lines.append("## 需要改进的地方")
        for i, imp in enumerate(report.improvements, 1):
            lines.append(f"{i}. {imp}")
        lines.append("")

    if report.corrected_scripts:
        lines.append("## 改进话术对照")
        for cs in report.corrected_scripts:
            lines.append(f"- **原话**：{cs.get('original', '')}")
            lines.append(f"  **改为**：{cs.get('corrected', '')}")
            lines.append(f"  **原因**：{cs.get('reason', '')}\n")

    if report.next_steps:
        lines.append("## 下一步行动建议")
        for i, ns in enumerate(report.next_steps, 1):
            lines.append(f"{i}. {ns}")
        lines.append("")

    if chat_history:
        lines.append("---\n## 话术讨论记录\n")
        for msg in chat_history:
            role = msg.get("role", "")
            content = msg.get("content", "")
            if isinstance(content, list):
                content = "".join(p.get("text", "") for p in content if isinstance(p, dict))
            if not content:
                continue
            label = "🤖 销售大师" if role == "assistant" else "👤 我"
            lines.append(f"**{label}**：{content}\n")

    return "\n".join(lines)


def format_evaluation_markdown(report) -> str:
    """将评估报告格式化为 Markdown 文本。"""
    lines = [f"# 训练评估报告\n"]
    lines.append(f"**综合评分**：{report.overall_score}/10\n")

    if report.dimension_scores:
        lines.append("## 各维度评分\n")
        for dim, data in report.dimension_scores.items():
            score = data.get("score", 0)
            bar = "●" * score + "○" * (10 - score)
            lines.append(f"**{dim}**：{bar} {score}/10")
            lines.append(f"  {data.get('justification', '')}\n")

    if report.strengths:
        lines.append("## 优势总结")
        for i, s in enumerate(report.strengths, 1):
            lines.append(f"{i}. {s}")
        lines.append("")

    if report.improvements:
        lines.append("## 改进建议")
        for i, s in enumerate(report.improvements, 1):
            lines.append(f"{i}. {s}")
        lines.append("")

    if report.style_alignment:
        sa = report.style_alignment
        lines.append("## 风格契合度")
        if sa.get("alignment_score") is not None:
            lines.append(f"契合度评分：{sa['alignment_score']}/10")
        if sa.get("matched_traits"):
            lines.append(f"契合特征：{', '.join(sa['matched_traits'])}")
        if sa.get("missed_traits"):
            lines.append(f"缺失特征：{', '.join(sa['missed_traits'])}")
        lines.append("")

    if report.conversation_summary:
        lines.append("## 实战总结\n")
        lines.append(report.conversation_summary)
        lines.append("")

    dp = report.deal_progression
    if dp:
        lines.append("## 签单路径分析\n")
        if dp.get("current_stage"):
            lines.append(f"**当前阶段**：{dp['current_stage']}")
        if dp.get("stage_progress"):
            progress_pct = int(dp["stage_progress"] * 100)
            lines.append(f"**推进进度**：{progress_pct}%")
        if dp.get("risk_level"):
            lines.append(f"**风险等级**：{dp['risk_level']}")
        if dp.get("blocking_issues"):
            lines.append("\n### 当前阻碍")
            for issue in dp["blocking_issues"]:
                lines.append(f"- {issue}")
        if dp.get("next_steps"):
            lines.append("\n### 下一步行动计划")
            for step in dp["next_steps"]:
                lines.append(f"\n**第{step.get('step', '?')}步：{step.get('action', '')}**")
                if step.get("script"):
                    lines.append(f"- 建议话术：{step['script']}")
                if step.get("goal"):
                    lines.append(f"- 目标：{step['goal']}")
        if dp.get("win_strategy"):
            lines.append(f"\n### 赢单策略\n{dp['win_strategy']}")

    return "\n".join(lines)


def export_as_docx(md_text: str, title: str = "销售大师报告") -> str:
    """将 Markdown 文本导出为 Word 文件，返回文件路径。"""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 40)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s*", "", stripped), style="List Number")
        else:
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            p = doc.add_paragraph(clean)

    path = os.path.join(tempfile.gettempdir(), f"{title}_{uuid.uuid4().hex[:6]}.docx")
    doc.save(path)
    return path


def generate_image(md_text: str, title: str = "销售大师") -> str:
    """将 Markdown 文本渲染为图片，返回文件路径。"""
    font_path = _find_chinese_font()
    fonts = {
        "title": ImageFont.truetype(font_path, 34),
        "subtitle": ImageFont.truetype(font_path, 20),
        "section": ImageFont.truetype(font_path, 25),
        "body": ImageFont.truetype(font_path, 20),
        "small": ImageFont.truetype(font_path, 15),
        "label": ImageFont.truetype(font_path, 16),
    }

    report_title, meta_lines, sections = _parse_markdown_for_poster(md_text)
    if not report_title:
        report_title = title
    if not sections and not meta_lines:
        sections = [{"title": "汇报内容", "lines": ["（无内容）"]}]
    is_evaluation = "训练评估" in title or "训练评估" in report_title
    hero_label = "训练评估" if is_evaluation else "面聊复盘"
    hero_subtitle = "Training Review" if is_evaluation else "Conversation Review"

    width = 1080
    card_width = 760
    panel_x = (width - card_width) // 2
    content_x = panel_x + 64
    content_w = card_width - 128
    top_margin = 120
    header_h = 72
    section_gap = 34

    measure = ImageDraw.Draw(Image.new("RGB", (1, 1)))
    body_lh = 34
    section_title_lh = 42
    total_content_h = header_h + 80
    total_content_h += _wrapped_height(report_title, content_w, measure, fonts["title"], 44)
    total_content_h += max(40, len(meta_lines) * 28)

    rendered_sections = []
    for idx, section in enumerate(sections):
        max_lines = 8 if idx < 3 else 6
        body_lines = []
        for raw in section["lines"]:
            cleaned = _clean_markdown_line(raw)
            if cleaned:
                body_lines.append(cleaned)
        if len(body_lines) > max_lines:
            body_lines = body_lines[:max_lines] + ["……"]
        section_h = 34 + section_title_lh
        for line in body_lines:
            section_h += _wrapped_height(line, content_w - 34, measure, fonts["body"], body_lh)
        section_h += 26
        rendered_sections.append((section["title"], body_lines, section_h))
        total_content_h += section_h + section_gap

    panel_h = total_content_h + 80
    height = max(1500, top_margin * 2 + panel_h)

    img = _make_textured_background(width, height)
    draw = ImageDraw.Draw(img)

    # Side marks inspired by editorial posters.
    draw.text((56, 90), "AI SALES MASTER", fill="#F5F7F8", font=fonts["label"])

    panel_y = top_margin
    draw.rounded_rectangle(
        [(panel_x, panel_y), (panel_x + card_width, panel_y + panel_h)],
        radius=0,
        fill="#F4FBFC",
    )

    # Compact report header.
    hero_y = panel_y
    draw.rectangle([(panel_x, hero_y), (panel_x + card_width, hero_y + header_h)], fill="#DDEEF1")
    draw.text((content_x, hero_y + 22), hero_label, fill="#235375", font=fonts["subtitle"])
    draw.text((content_x + 110, hero_y + 27), hero_subtitle, fill="#6E8795", font=fonts["small"])

    y = panel_y + header_h + 54
    y = _draw_wrapped_text(draw, report_title, content_x, y, content_w, fonts["title"], "#183B56", 44)
    y += 12
    for meta in meta_lines[:2]:
        y = _draw_wrapped_text(draw, _clean_markdown_line(meta), content_x, y, content_w, fonts["small"], "#587284", 25)
    y += 28

    accent_colors = ["#255D8A", "#7A9E9F", "#B58B4B", "#556B8E", "#8B6F7E"]
    for idx, (section_title, body_lines, section_h) in enumerate(rendered_sections):
        color = accent_colors[idx % len(accent_colors)]
        draw.rectangle([(content_x, y + 8), (content_x + 14, y + 42)], fill=color)
        draw.text((content_x + 28, y), section_title, fill=color, font=fonts["section"])
        y += section_title_lh + 12
        for line in body_lines:
            y = _draw_wrapped_text(draw, line, content_x + 28, y, content_w - 34, fonts["body"], "#2E4656", body_lh)
        y += 24
        if idx != len(rendered_sections) - 1:
            draw.line([(content_x, y), (content_x + content_w, y)], fill="#D5E2E6", width=1)
            y += section_gap

    footer_y = panel_y + panel_h - 54
    draw.line([(content_x, footer_y - 16), (content_x + content_w, footer_y - 16)], fill="#D5E2E6", width=1)
    draw.text((content_x, footer_y), "AI 销售大师 · 克拉时刻", fill="#7F949F", font=fonts["small"])

    path = os.path.join(tempfile.gettempdir(), f"{title}_{uuid.uuid4().hex[:6]}.png")
    img.save(path, "PNG")
    return path


def generate_share_link(md_text: str, title: str = "销售大师报告") -> str:
    """生成一个本地可访问的分享链接，返回 URL。"""
    html_content = _build_share_html(md_text, title)

    share_id = uuid.uuid4().hex[:8]
    share_dir = os.path.join(tempfile.gettempdir(), "sales_master_shares")
    os.makedirs(share_dir, exist_ok=True)

    share_file = os.path.join(share_dir, f"{share_id}.html")
    with open(share_file, "w", encoding="utf-8") as f:
        f.write(html_content)

    # Try to find an available port
    port = _find_available_port(7870, 7890)
    if port and not _share_server_running(port):
        _start_share_server(share_dir, port)

    return f"http://localhost:{port}/{share_id}.html"


def _wrap_text(text, max_width, draw, font):
    """Wrap text to fit within max_width pixels."""
    lines = []
    current = ""
    for char in text:
        test = current + char
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] - bbox[0] > max_width:
            if current:
                lines.append(current)
            current = char
        else:
            current = test
    if current:
        lines.append(current)
    return lines if lines else [text]


def _find_chinese_font() -> str:
    font_paths = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/msyhbd.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    ]
    for fp in font_paths:
        if os.path.exists(fp) and _font_can_render_chinese(fp):
            return fp
    raise RuntimeError(
        "未找到可用中文字体，无法生成中文图片。请在服务器安装 fonts-noto-cjk 或 fonts-wqy-zenhei。"
    )


def _font_can_render_chinese(font_path: str) -> bool:
    try:
        font = ImageFont.truetype(font_path, 20)
        return font.getmask("销售汇报").getbbox() is not None
    except Exception:
        return False


def _parse_markdown_for_poster(md_text: str):
    title = ""
    meta_lines = []
    sections = []
    current = None

    for raw_line in md_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            title = _clean_markdown_line(line[2:])
            continue
        if line.startswith("## "):
            if current:
                sections.append(current)
            current = {"title": _clean_markdown_line(line[3:]), "lines": []}
            continue
        if line.startswith("### "):
            if current:
                sections.append(current)
            current = {"title": _clean_markdown_line(line[4:]), "lines": []}
            continue
        if line.startswith("---"):
            continue
        if current is None:
            meta_lines.append(line)
        else:
            current["lines"].append(line)

    if current:
        sections.append(current)
    return title, meta_lines, sections


def _clean_markdown_line(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"[*_](.*?)[*_]", r"\1", text)
    text = re.sub(r"^\s*[-•]\s*", "• ", text)
    text = re.sub(r"^\s*(\d+)\.\s*", r"\1. ", text)
    return text.strip()


def _wrapped_height(text: str, max_width: int, draw, font, line_height: int) -> int:
    return max(line_height, len(_wrap_text(text, max_width, draw, font)) * line_height)


def _draw_wrapped_text(draw, text: str, x: int, y: int, max_width: int, font, fill: str, line_height: int) -> int:
    for line in _wrap_text(text, max_width, draw, font):
        draw.text((x, y), line, fill=fill, font=font)
        y += line_height
    return y


def _make_textured_background(width: int, height: int) -> Image.Image:
    base = Image.new("RGB", (width, height), "#596A76")
    noise = Image.effect_noise((width, height), 38).convert("L")
    texture = Image.new("RGB", (width, height), "#6E7E87")
    base = Image.blend(base, texture, 0.38)
    base.putalpha(255)
    noise_rgb = Image.merge("RGB", (noise, noise, noise))
    return Image.blend(base.convert("RGB"), noise_rgb, 0.08)


def _build_share_html(md_text: str, title: str) -> str:
    """Build a self-contained HTML page for sharing."""
    import html as html_module
    escaped = html_module.escape(md_text)

    # Simple markdown-to-HTML conversion
    body = escaped
    body = re.sub(r"^### (.+)$", r"<h3>\1</h3>", body, flags=re.MULTILINE)
    body = re.sub(r"^## (.+)$", r"<h2>\1</h2>", body, flags=re.MULTILINE)
    body = re.sub(r"^# (.+)$", r"<h1>\1</h1>", body, flags=re.MULTILINE)
    body = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", body)
    body = re.sub(r"^---$", "<hr>", body, flags=re.MULTILINE)
    body = re.sub(r"^- (.+)$", r"<li>\1</li>", body, flags=re.MULTILINE)
    body = re.sub(r"^(\d+)\. (.+)$", r"<li>\2</li>", body, flags=re.MULTILINE)
    body = re.sub(r"\n\n", "<br><br>", body)
    body = re.sub(r"\n", "<br>", body)

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{html_module.escape(title)}</title>
<style>
body {{ font-family: -apple-system, "Microsoft YaHei", sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; color: #333; line-height: 1.8; }}
h1 {{ color: #2e7d32; border-bottom: 2px solid #4CAF50; padding-bottom: 8px; }}
h2 {{ color: #388e3c; margin-top: 24px; }}
h3 {{ color: #43a047; }}
strong {{ color: #1b5e20; }}
hr {{ border: none; border-top: 1px solid #e0e0e0; margin: 20px 0; }}
li {{ margin: 4px 0; }}
.footer {{ text-align: center; color: #999; font-size: 12px; margin-top: 30px; padding-top: 10px; border-top: 1px solid #eee; }}
</style>
</head>
<body>
{body}
<div class="footer">AI 销售大师 · 克拉时刻</div>
</body>
</html>"""


_share_server_instance = None
_share_server_port = None


def _find_available_port(start, end):
    import socket
    for port in range(start, end + 1):
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(1)
            result = s.connect_ex(("localhost", port))
            s.close()
            if result != 0:
                return port
        except Exception:
            continue
    return None


def _share_server_running(port):
    import socket
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.settimeout(1)
        result = s.connect_ex(("localhost", port))
        s.close()
        return result == 0
    except Exception:
        return False


class _ShareHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        share_dir = os.path.join(tempfile.gettempdir(), "sales_master_shares")
        filename = self.path.lstrip("/")
        filepath = os.path.join(share_dir, filename)
        if os.path.exists(filepath) and filename.endswith(".html"):
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            with open(filepath, "rb") as f:
                content = f.read()
            self.send_header("Content-Length", len(content))
            self.end_headers()
            self.wfile.write(content)
        else:
            self.send_response(404)
            self.end_headers()
            self.wfile.write(b"Not Found")

    def log_message(self, format, *args):
        pass  # Suppress logging


def _start_share_server(share_dir, port):
    global _share_server_instance, _share_server_port
    _share_server_port = port
    server = HTTPServer(("0.0.0.0", port), _ShareHandler)
    _share_server_instance = server
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
